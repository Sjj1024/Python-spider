from docx import *
from docx.shared import Inches

# 创建一个doc文档
file = Document()

# 添加一个标题，等级为2,也就是2级标题
file.add_heading("段落标题", level=2)

# 添加段落内容
para = file.add_paragraph("我好喜欢你")

# 添加图片，图片是当前文件夹下的 img.png 图片
file.add_picture('06b01.jpg',width=Inches(4.0))

# 保存文件
file.save("test.docx")
